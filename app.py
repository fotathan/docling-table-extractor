import io
import json
import zipfile
from pathlib import Path
from datetime import datetime
from typing import Any

import pandas as pd
import streamlit as st


# --------------------------------------------------
# App config
# --------------------------------------------------
st.set_page_config(
    page_title="PDF/XLSX/DOCX Table Extractor",
    page_icon="📄",
    layout="wide"
)

BASE_DIR = Path(".")
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# --------------------------------------------------
# UI header
# --------------------------------------------------
st.title("📄 PDF / XLSX / DOCX Table Extractor")
st.caption("Upload PDF, XLSX, or DOCX files and extract text previews and tables.")

with st.container():
    st.markdown(
        """
        **This MVP does:**
        - PDF via Camelot + PyPDF
        - XLSX via pandas
        - DOCX text + Word tables via python-docx
        - table previews
        - CSV / HTML / Markdown / JSON exports
        - ZIP download of all generated outputs

        **Cloud-safe direction**
        - no Docling
        - no OCR
        - best for text-based PDFs
        """
    )


# --------------------------------------------------
# Helpers
# --------------------------------------------------
def safe_filename(name: str) -> str:
    keepchars = (" ", ".", "_", "-")
    cleaned = "".join(c for c in name if c.isalnum() or c in keepchars).strip()
    return cleaned or "uploaded_file"


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def save_uploaded_file(uploaded_file) -> Path:
    filename = f"{now_stamp()}_{safe_filename(uploaded_file.name)}"
    file_path = UPLOAD_DIR / filename
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path


def write_text_file(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def write_json_file(path: Path, data: Any) -> None:
    path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8"
    )


def build_zip(file_paths: list[Path]) -> io.BytesIO:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_path in file_paths:
            if file_path.exists() and file_path.is_file():
                zf.write(file_path, arcname=file_path.name)
    zip_buffer.seek(0)
    return zip_buffer


def get_job_output_dir(file_path: Path) -> Path:
    job_dir = OUTPUT_DIR / file_path.stem
    job_dir.mkdir(parents=True, exist_ok=True)
    return job_dir


def dataframe_to_markdown_fallback(df: pd.DataFrame, title: str) -> str:
    lines = [f"### {title}", ""]
    if df.empty:
        lines.append("_Empty table_")
        return "\n".join(lines)

    header = "| " + " | ".join(str(c) for c in df.columns) + " |"
    sep = "| " + " | ".join(["---"] * len(df.columns)) + " |"
    lines.extend([header, sep])

    max_rows = min(len(df), 20)
    for _, row in df.head(max_rows).iterrows():
        lines.append("| " + " | ".join("" if pd.isna(v) else str(v) for v in row.tolist()) + " |")

    if len(df) > max_rows:
        lines.append("")
        lines.append(f"_Preview truncated. Total rows: {len(df)}_")

    return "\n".join(lines)


def html_table_from_df(df: pd.DataFrame) -> str:
    return df.to_html(index=False, border=0)


# --------------------------------------------------
# Availability checks
# --------------------------------------------------
def check_pdf_stack_available() -> tuple[bool, str]:
    try:
        import camelot  # noqa: F401
        from pypdf import PdfReader  # noqa: F401
        return True, ""
    except Exception as e:
        return False, str(e)


def check_docx_stack_available() -> tuple[bool, str]:
    try:
        import docx  # noqa: F401
        return True, ""
    except Exception as e:
        return False, str(e)


# --------------------------------------------------
# XLSX processing
# --------------------------------------------------
def process_xlsx_with_pandas(file_path: Path, job_output_dir: Path) -> dict:
    result = {
        "file_name": file_path.name,
        "file_type": "xlsx",
        "mode": "xlsx_pandas",
        "status": "success",
        "preview_text": "",
        "tables": [],
        "output_files": [],
        "errors": [],
    }

    try:
        excel_file = pd.ExcelFile(file_path)
        preview_lines = [f"# Spreadsheet: {file_path.name}", ""]
        preview_lines.append(f"Sheets detected: {', '.join(excel_file.sheet_names)}")
        preview_lines.append("")

        workbook_summary = {
            "file": file_path.name,
            "mode": result["mode"],
            "sheet_count": len(excel_file.sheet_names),
            "sheets": [],
            "errors": [],
        }

        for idx, sheet_name in enumerate(excel_file.sheet_names, start=1):
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                csv_path = job_output_dir / f"{file_path.stem}_sheet_{idx}_{safe_filename(sheet_name)}.csv"
                html_path = job_output_dir / f"{file_path.stem}_sheet_{idx}_{safe_filename(sheet_name)}.html"

                df.to_csv(csv_path, index=False)
                write_text_file(html_path, html_table_from_df(df))

                table_info = {
                    "index": idx,
                    "name": sheet_name,
                    "dataframe": df,
                    "csv_path": csv_path,
                    "html_path": html_path,
                    "markdown": dataframe_to_markdown_fallback(df, f"Sheet {idx}: {sheet_name}"),
                    "rows": len(df),
                    "columns": len(df.columns),
                }

                result["tables"].append(table_info)
                result["output_files"].extend([csv_path, html_path])

                workbook_summary["sheets"].append({
                    "name": sheet_name,
                    "rows": len(df),
                    "columns": list(df.columns),
                    "csv_file": csv_path.name,
                })

                preview_lines.append(f"## Sheet {idx}: {sheet_name}")
                preview_lines.append(f"Rows: {len(df)}")
                preview_lines.append(f"Columns: {len(df.columns)}")
                preview_lines.append("")

            except Exception as e:
                msg = f"Failed reading sheet '{sheet_name}': {e}"
                result["errors"].append(msg)
                workbook_summary["errors"].append(msg)

        preview_text = "\n".join(preview_lines)
        result["preview_text"] = preview_text

        md_path = job_output_dir / f"{file_path.stem}_summary.md"
        json_path = job_output_dir / f"{file_path.stem}_summary.json"

        write_text_file(md_path, preview_text)
        write_json_file(json_path, workbook_summary)

        result["output_files"].extend([md_path, json_path])

        if result["errors"] and not result["tables"]:
            result["status"] = "failed"
        elif result["errors"]:
            result["status"] = "partial_success"

    except Exception as e:
        result["status"] = "failed"
        result["errors"].append(f"XLSX processing failed: {e}")

    return result


# --------------------------------------------------
# DOCX processing
# --------------------------------------------------
def process_docx(file_path: Path, job_output_dir: Path) -> dict:
    result = {
        "file_name": file_path.name,
        "file_type": "docx",
        "mode": "python_docx",
        "status": "success",
        "preview_text": "",
        "tables": [],
        "output_files": [],
        "errors": [],
    }

    ok, err = check_docx_stack_available()
    if not ok:
        result["status"] = "failed"
        result["errors"].append(f"DOCX stack import failed: {err}")
        return result

    try:
        import docx

        doc = docx.Document(str(file_path))

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        preview_text = "\n".join(paragraphs[:80])
        result["preview_text"] = preview_text

        md_path = job_output_dir / f"{file_path.stem}.md"
        json_path = job_output_dir / f"{file_path.stem}.json"

        write_text_file(md_path, preview_text)
        result["output_files"].append(md_path)

        doc_summary = {
            "file": file_path.name,
            "mode": result["mode"],
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "errors": [],
        }

        for idx, table in enumerate(doc.tables, start=1):
            try:
                rows = []
                max_cols = 0

                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    rows.append(values)
                    max_cols = max(max_cols, len(values))

                normalized_rows = []
                for row in rows:
                    normalized_rows.append(row + [""] * (max_cols - len(row)))

                if normalized_rows:
                    header = normalized_rows[0]
                    data_rows = normalized_rows[1:] if len(normalized_rows) > 1 else []
                    if any(str(h).strip() for h in header):
                        df = pd.DataFrame(data_rows, columns=header)
                    else:
                        df = pd.DataFrame(normalized_rows)
                else:
                    df = pd.DataFrame()

                csv_path = job_output_dir / f"{file_path.stem}_table_{idx}.csv"
                html_path = job_output_dir / f"{file_path.stem}_table_{idx}.html"

                df.to_csv(csv_path, index=False)
                write_text_file(html_path, html_table_from_df(df))

                table_info = {
                    "index": idx,
                    "name": f"Table {idx}",
                    "dataframe": df,
                    "csv_path": csv_path,
                    "html_path": html_path,
                    "markdown": dataframe_to_markdown_fallback(df, f"Table {idx}"),
                    "rows": len(df),
                    "columns": len(df.columns),
                }

                result["tables"].append(table_info)
                result["output_files"].extend([csv_path, html_path])

            except Exception as e:
                msg = f"DOCX table {idx} failed: {e}"
                result["errors"].append(msg)
                doc_summary["errors"].append(msg)

        write_json_file(json_path, doc_summary)
        result["output_files"].append(json_path)

        if result["errors"] and not result["preview_text"] and not result["tables"]:
            result["status"] = "failed"
        elif result["errors"]:
            result["status"] = "partial_success"

    except Exception as e:
        result["status"] = "failed"
        result["errors"].append(f"DOCX processing failed: {e}")

    return result


# --------------------------------------------------
# PDF processing with Camelot
# --------------------------------------------------
def extract_pdf_text(file_path: Path) -> tuple[str, list[str]]:
    errors = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        parts = []
        for page in reader.pages[:25]:
            try:
                parts.append(page.extract_text() or "")
            except Exception as e:
                errors.append(f"PDF text extraction warning: {e}")

        return "\n\n".join([p.strip() for p in parts if p and p.strip()]), errors
    except Exception as e:
        return "", [f"PDF text extraction failed: {e}"]


def run_camelot(file_path: Path, flavor: str):
    import camelot
    return camelot.read_pdf(str(file_path), pages="all", flavor=flavor)


def process_pdf_with_camelot(file_path: Path, job_output_dir: Path) -> dict:
    result = {
        "file_name": file_path.name,
        "file_type": "pdf",
        "mode": "camelot",
        "status": "success",
        "preview_text": "",
        "tables": [],
        "output_files": [],
        "errors": [],
    }

    ok, err = check_pdf_stack_available()
    if not ok:
        result["status"] = "failed"
        result["errors"].append(f"PDF stack import failed: {err}")
        return result

    preview_text, text_errors = extract_pdf_text(file_path)
    result["preview_text"] = preview_text
    result["errors"].extend(text_errors)

    md_path = job_output_dir / f"{file_path.stem}.md"
    write_text_file(md_path, preview_text or "No text preview available.")
    result["output_files"].append(md_path)

    extraction_summary = {
        "file": file_path.name,
        "mode": result["mode"],
        "flavor_used": None,
        "table_count": 0,
        "errors": [],
    }

    tables_obj = None
    used_flavor = None

    # Try lattice first, then stream fallback
    for flavor in ["lattice", "stream"]:
        try:
            candidate = run_camelot(file_path, flavor)
            if len(candidate) > 0:
                tables_obj = candidate
                used_flavor = flavor
                break
        except Exception as e:
            msg = f"Camelot {flavor} failed: {e}"
            result["errors"].append(msg)
            extraction_summary["errors"].append(msg)

    # If lattice found 0 tables but no exception, try stream too
    if tables_obj is None:
        try:
            candidate = run_camelot(file_path, "stream")
            tables_obj = candidate
            used_flavor = "stream"
        except Exception as e:
            msg = f"Camelot stream fallback failed: {e}"
            result["errors"].append(msg)
            extraction_summary["errors"].append(msg)

    if tables_obj is not None:
        extraction_summary["flavor_used"] = used_flavor
        extraction_summary["table_count"] = len(tables_obj)

        for idx, table in enumerate(tables_obj, start=1):
            try:
                df = table.df.copy()

                # Try to use first row as header when it looks sensible
                if not df.empty and len(df) > 1:
                    first_row = [str(x).strip() for x in df.iloc[0].tolist()]
                    unique_nonempty = len(set([x for x in first_row if x]))
                    if unique_nonempty >= 1:
                        df.columns = first_row
                        df = df.iloc[1:].reset_index(drop=True)

                csv_path = job_output_dir / f"{file_path.stem}_table_{idx}.csv"
                html_path = job_output_dir / f"{file_path.stem}_table_{idx}.html"

                df.to_csv(csv_path, index=False)
                write_text_file(html_path, html_table_from_df(df))

                table_info = {
                    "index": idx,
                    "name": f"Table {idx} ({used_flavor})",
                    "dataframe": df,
                    "csv_path": csv_path,
                    "html_path": html_path,
                    "markdown": dataframe_to_markdown_fallback(df, f"Table {idx}"),
                    "rows": len(df),
                    "columns": len(df.columns),
                }

                result["tables"].append(table_info)
                result["output_files"].extend([csv_path, html_path])

            except Exception as e:
                msg = f"PDF table {idx} failed: {e}"
                result["errors"].append(msg)
                extraction_summary["errors"].append(msg)

    json_path = job_output_dir / f"{file_path.stem}.json"
    write_json_file(json_path, extraction_summary)
    result["output_files"].append(json_path)

    if result["errors"] and not result["preview_text"] and not result["tables"]:
        result["status"] = "failed"
    elif result["errors"]:
        result["status"] = "partial_success"

    return result


# --------------------------------------------------
# Router
# --------------------------------------------------
def process_file(file_path: Path):
    job_output_dir = get_job_output_dir(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".xlsx":
        return process_xlsx_with_pandas(file_path, job_output_dir)

    if suffix == ".pdf":
        return process_pdf_with_camelot(file_path, job_output_dir)

    if suffix == ".docx":
        return process_docx(file_path, job_output_dir)

    return {
        "file_name": file_path.name,
        "file_type": suffix.replace(".", ""),
        "mode": "unsupported",
        "status": "failed",
        "preview_text": "",
        "tables": [],
        "output_files": [],
        "errors": [f"Unsupported file type: {suffix}"],
    }


# --------------------------------------------------
# Sidebar
# --------------------------------------------------
with st.sidebar:
    st.header("Options")
    show_full_preview = st.checkbox("Show full extracted preview", value=False)
    max_preview_chars = st.slider("Preview length", min_value=1000, max_value=15000, value=4000, step=500)
    st.markdown("---")
    st.info("Tip: Start with 3-5 real files from your own use case.")
    st.caption("PDFs use Camelot for tables and PyPDF for text preview.")

    pdf_ok, pdf_err = check_pdf_stack_available()
    if pdf_ok:
        st.success("PDF stack import check passed")
    else:
        st.error("PDF stack import check failed")
        st.code(pdf_err)

    docx_ok, docx_err = check_docx_stack_available()
    if docx_ok:
        st.success("DOCX stack import check passed")
    else:
        st.error("DOCX stack import check failed")
        st.code(docx_err)


# --------------------------------------------------
# Upload
# --------------------------------------------------
uploaded_files = st.file_uploader(
    "Upload one or more files",
    type=["pdf", "docx", "xlsx"],
    accept_multiple_files=True,
    help="Supported: PDF, DOCX, XLSX"
)

process_clicked = st.button("🚀 Process uploaded files", type="primary", disabled=not uploaded_files)


# --------------------------------------------------
# Processing
# --------------------------------------------------
if process_clicked and uploaded_files:
    saved_paths = []
    all_results = []
    all_output_files = []

    progress_bar = st.progress(0, text="Preparing files...")
    status_placeholder = st.empty()

    total_files = len(uploaded_files)

    for file in uploaded_files:
        saved_paths.append(save_uploaded_file(file))

    for idx, file_path in enumerate(saved_paths, start=1):
        status_placeholder.info(f"Processing {idx}/{total_files}: {file_path.name}")

        file_result = process_file(file_path)
        all_results.append(file_result)
        all_output_files.extend(file_result["output_files"])

        progress_bar.progress(idx / total_files, text=f"Processed {idx}/{total_files} files")

    status_placeholder.success("Processing completed.")

    total_tables = sum(len(r["tables"]) for r in all_results)
    success_count = sum(1 for r in all_results if r["status"] == "success")
    partial_count = sum(1 for r in all_results if r["status"] == "partial_success")
    failed_count = sum(1 for r in all_results if r["status"] == "failed")

    st.markdown("## Results Overview")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Files processed", len(all_results))
    c2.metric("Tables found", total_tables)
    c3.metric("Successful", success_count)
    c4.metric("Partial / Failed", partial_count + failed_count)

    summary_tab, preview_tab, tables_tab, downloads_tab = st.tabs(
        ["Summary", "Preview", "Tables", "Downloads"]
    )

    with summary_tab:
        summary_rows = []
        for r in all_results:
            summary_rows.append({
                "File": r["file_name"],
                "Type": r["file_type"],
                "Mode": r["mode"],
                "Status": r["status"],
                "Tables": len(r["tables"]),
                "Errors": len(r["errors"]),
            })

        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True)

        for r in all_results:
            with st.expander(f"{r['file_name']} ({r['status']})"):
                if r["errors"]:
                    for err in r["errors"]:
                        st.warning(err)
                else:
                    st.success("No issues reported.")

    with preview_tab:
        for r in all_results:
            with st.expander(f"Preview - {r['file_name']}", expanded=False):
                preview = r["preview_text"] or "No preview available."
                if not show_full_preview and len(preview) > max_preview_chars:
                    preview = preview[:max_preview_chars] + "\n\n...[truncated]"

                st.text_area(
                    label=f"Extracted preview for {r['file_name']}",
                    value=preview,
                    height=280,
                    key=f"preview_{r['file_name']}"
                )

    with tables_tab:
        any_tables = any(r["tables"] for r in all_results)

        if not any_tables:
            st.info("No tables detected in the uploaded files.")
        else:
            for r in all_results:
                if not r["tables"]:
                    continue

                st.markdown(f"### {r['file_name']}")

                for table in r["tables"]:
                    table_title = table.get("name") or f"Table {table.get('index', '?')}"
                    with st.expander(table_title, expanded=False):
                        if table.get("markdown"):
                            st.markdown(table["markdown"])

                        df = table.get("dataframe")
                        if df is not None:
                            st.dataframe(df, use_container_width=True)

                        meta_cols = st.columns(2)
                        meta_cols[0].write(f"Rows: {table.get('rows', '-')}")
                        meta_cols[1].write(f"Columns: {table.get('columns', '-')}")

                        csv_path = table.get("csv_path")
                        if csv_path and Path(csv_path).exists():
                            with open(csv_path, "rb") as f:
                                st.download_button(
                                    label="Download CSV",
                                    data=f,
                                    file_name=Path(csv_path).name,
                                    mime="text/csv",
                                    key=f"csv_{r['file_name']}_{table['index']}"
                                )

                        html_path = table.get("html_path")
                        if html_path and Path(html_path).exists():
                            with open(html_path, "rb") as f:
                                st.download_button(
                                    label="Download HTML",
                                    data=f,
                                    file_name=Path(html_path).name,
                                    mime="text/html",
                                    key=f"html_{r['file_name']}_{table['index']}"
                                )

    with downloads_tab:
        st.subheader("Download everything")

        if all_output_files:
            zip_buffer = build_zip(all_output_files)
            st.download_button(
                label="📦 Download all generated outputs as ZIP",
                data=zip_buffer,
                file_name=f"table_extractor_outputs_{now_stamp()}.zip",
                mime="application/zip"
            )
        else:
            st.info("No generated outputs available.")

        download_rows = []
        for r in all_results:
            for p in r["output_files"]:
                download_rows.append({
                    "Source file": r["file_name"],
                    "Generated file": Path(p).name,
                    "Path": str(p),
                })

        if download_rows:
            st.dataframe(pd.DataFrame(download_rows), use_container_width=True)

else:
    st.info("Upload one or more files and click **Process uploaded files**.")
